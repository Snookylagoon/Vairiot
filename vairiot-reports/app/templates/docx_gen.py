"""
DOCX report generator — python-docx with full Vairiot livery.

Layout (A4):
  - Top gradient accent bar (table-based, pink → mauve → violet)
  - Letterhead: VAIRIOT brand left, company info right
  - Report title + subtitle
  - Metadata line (filters, date)
  - Summary metrics (bordered-left cards)
  - Data table with repeating headers and banded rows
  - Authorisation / signoff block
  - Footer: confidentiality + page numbers
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any

from docx import Document
from docx.shared import Mm, Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from app.config import (
    BRAND, ColType, ReportDef,
    A4_WIDTH_MM, A4_HEIGHT_MM,
    MARGIN_TOP_MM, MARGIN_BOTTOM_MM, MARGIN_LEFT_MM, MARGIN_RIGHT_MM,
)
from app.models import ReportRequest
from app.templates.brand import format_value, hex_to_rgb, interpolate_gradient


# ── Colour helpers ────────────────────────────────────────────────────────────

def _rgb(hex_colour: str) -> RGBColor:
    r, g, b = hex_to_rgb(hex_colour)
    return RGBColor(r, g, b)

_C_PINK     = _rgb(BRAND.pink)
_C_MAUVE    = _rgb(BRAND.mauve)
_C_VIOLET   = _rgb(BRAND.violet)
_C_CHARCOAL = _rgb(BRAND.charcoal)
_C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
_C_WASH     = _rgb(BRAND.wash)
_C_GREY     = RGBColor(0x88, 0x88, 0x88)
_C_GREEN    = _rgb(BRAND.green)
_C_RED      = _rgb(BRAND.red)
_C_AMBER    = _rgb(BRAND.amber)


def _set_cell_shading(cell, hex_colour: str):
    """Set cell background fill colour."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_colour.lstrip("#")}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border_left(cell, hex_colour: str, width_pt: float = 2.0):
    """Set a left border on a cell (for metric cards)."""
    width_eighths = int(width_pt * 8)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{width_eighths}" w:space="0" w:color="{hex_colour.lstrip("#")}"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)


def _set_repeat_header_row(row):
    """Mark a table row to repeat on every page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>')
    trPr.append(tbl_header)


def _set_row_height(row, mm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(mm * 56.7)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def _add_run(paragraph, text: str, font_name: str = "Montserrat",
             size_pt: float = 9, bold: bool = False,
             colour: RGBColor = _C_CHARCOAL):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = colour
    return run


def _col_align(col_type: str) -> int:
    if col_type in (ColType.CURRENCY, ColType.NUMBER, ColType.INTEGER, ColType.PERCENT):
        return WD_ALIGN_PARAGRAPH.RIGHT
    if col_type == ColType.DATE:
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def generate_docx(report_def: ReportDef, req: ReportRequest) -> io.BytesIO:
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────
    section = doc.sections[0]
    if report_def.is_landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(A4_HEIGHT_MM)
        section.page_height = Mm(A4_WIDTH_MM)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(A4_WIDTH_MM)
        section.page_height = Mm(A4_HEIGHT_MM)

    section.top_margin    = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin   = Mm(MARGIN_LEFT_MM)
    section.right_margin  = Mm(MARGIN_RIGHT_MM)

    # ── Gradient accent bar ───────────────────────────────────────────────
    gradient_steps = 30
    gradient_colours = interpolate_gradient(BRAND.gradient_stops(), gradient_steps)
    accent_table = doc.add_table(rows=1, cols=gradient_steps)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, colour in enumerate(gradient_colours):
        cell = accent_table.cell(0, i)
        _set_cell_shading(cell, colour)
        cell.text = ""
        cell.paragraphs[0].space_before = Pt(0)
        cell.paragraphs[0].space_after = Pt(0)
    _set_row_height(accent_table.rows[0], 2.0)
    _remove_table_borders(accent_table)

    # ── Letterhead ────────────────────────────────────────────────────────
    # Brand name: "VAIR" in charcoal, "IOT" per-letter gradient (pink→mauve→violet)
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(0)
    run_vair = p.add_run("VAIR")
    run_vair.font.name = "Montserrat"
    run_vair.font.size = Pt(22)
    run_vair.font.bold = True
    run_vair.font.color.rgb = _C_CHARCOAL

    iot_colours = [_C_PINK, _C_MAUVE, _C_VIOLET]
    for i, letter in enumerate("IOT"):
        run_l = p.add_run(letter)
        run_l.font.name = "Montserrat"
        run_l.font.size = Pt(22)
        run_l.font.bold = True
        run_l.font.color.rgb = iot_colours[i]

    # Tag line
    p_tag = doc.add_paragraph()
    p_tag.space_before = Pt(0)
    p_tag.space_after = Pt(4)
    _add_run(p_tag, "ASSET INTELLIGENCE", size_pt=7, colour=_C_GREY)

    # Company info (right-aligned block)
    if req.company.legal_name:
        p_co = doc.add_paragraph()
        p_co.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_co.space_before = Pt(0)
        p_co.space_after = Pt(2)
        _add_run(p_co, req.company.legal_name, size_pt=9, bold=True, colour=_C_CHARCOAL)
        if req.company.full_address:
            p_addr = doc.add_paragraph()
            p_addr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_addr.space_before = Pt(0)
            p_addr.space_after = Pt(0)
            _add_run(p_addr, req.company.full_address, size_pt=7, colour=_C_GREY)
        if req.company.registration_number:
            p_reg = doc.add_paragraph()
            p_reg.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_reg.space_before = Pt(0)
            p_reg.space_after = Pt(0)
            _add_run(p_reg, req.company.registration_number, size_pt=7, colour=_C_GREY)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.space_before = Pt(4)
    p_div.space_after = Pt(8)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="D9D9D9"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── Report title ──────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.space_before = Pt(4)
    p_title.space_after = Pt(2)
    _add_run(p_title, report_def.title, size_pt=16, bold=True, colour=_C_VIOLET)

    # Subtitle / metadata line
    meta_parts = []
    if report_def.subtitle:
        meta_parts.append(report_def.subtitle)
    for k, v in req.filters.items():
        label = k.replace("_", " ").replace("Id", "").title()
        meta_parts.append(f"{label}: {v}")
    meta_parts.append(f"Generated: {date.today().strftime('%d %B %Y')}")

    p_meta = doc.add_paragraph()
    p_meta.space_before = Pt(0)
    p_meta.space_after = Pt(8)
    _add_run(p_meta, " · ".join(meta_parts), size_pt=8, colour=_C_GREY)

    # ── Summary metrics ───────────────────────────────────────────────────
    if report_def.summary_fields and req.summary:
        num_metrics = len(report_def.summary_fields)
        summary_table = doc.add_table(rows=1, cols=num_metrics)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        accent_colours = [BRAND.green, BRAND.red, BRAND.amber, BRAND.violet,
                          BRAND.blue, BRAND.pink, BRAND.mauve]

        for i, sf in enumerate(report_def.summary_fields):
            cell = summary_table.cell(0, i)
            _set_cell_shading(cell, BRAND.wash)
            _set_cell_border_left(cell, accent_colours[i % len(accent_colours)])

            val = req.summary.get(sf["key"], "")
            formatted = format_value(val, sf.get("type", ColType.TEXT), req.currency)

            p_val = cell.paragraphs[0]
            p_val.space_before = Pt(4)
            p_val.space_after = Pt(0)
            _add_run(p_val, formatted, size_pt=14, bold=True,
                     colour=_rgb(accent_colours[i % len(accent_colours)]))

            p_lbl = cell.add_paragraph()
            p_lbl.space_before = Pt(0)
            p_lbl.space_after = Pt(4)
            _add_run(p_lbl, sf["label"], size_pt=7, colour=_C_GREY)

        _set_row_height(summary_table.rows[0], 16.0)

        # Remove inner borders from summary table
        _remove_table_borders(summary_table)

        # Spacer
        doc.add_paragraph().space_after = Pt(4)

    # ── Data table ────────────────────────────────────────────────────────
    num_cols = len(report_def.columns)
    num_data_rows = len(req.rows)
    total_rows = 1 + num_data_rows  # header + data
    if report_def.show_totals and req.totals:
        total_rows += 1

    table = doc.add_table(rows=total_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Set column widths
    col_widths_pt = report_def.resolve_column_widths_pt()
    for i, width_pt in enumerate(col_widths_pt):
        for row in table.rows:
            cell = row.cells[i]
            cell.width = Emu(int(width_pt * 914400 / 72))

    # Header row
    header_row = table.rows[0]
    _set_repeat_header_row(header_row)
    for i, col_def in enumerate(report_def.columns):
        cell = header_row.cells[i]
        _set_cell_shading(cell, BRAND.violet)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = _col_align(col_def.col_type)
        _add_run(p, col_def.header, size_pt=8, bold=True, colour=_C_WHITE)
    _set_row_height(header_row, 7.0)

    # Data rows
    for row_idx in range(num_data_rows):
        row = table.rows[1 + row_idx]
        row_data = req.rows[row_idx]
        is_alt = row_idx % 2 == 1

        for col_idx, col_def in enumerate(report_def.columns):
            cell = row.cells[col_idx]
            if is_alt:
                _set_cell_shading(cell, BRAND.wash)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = _col_align(col_def.col_type)

            raw_val = row_data.get(col_def.key)
            formatted = format_value(raw_val, col_def.col_type, req.currency)

            font_name = "Montserrat" if col_def.col_type in (
                ColType.CURRENCY, ColType.NUMBER, ColType.INTEGER
            ) else "Montserrat"
            _add_run(p, formatted, font_name=font_name, size_pt=7.5, colour=_C_CHARCOAL)

        _set_row_height(row, 6.0)

    # Totals row
    if report_def.show_totals and req.totals:
        totals_row = table.rows[-1]
        for col_idx, col_def in enumerate(report_def.columns):
            cell = totals_row.cells[col_idx]
            _set_cell_shading(cell, BRAND.violet)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = _col_align(col_def.col_type)

            if col_idx == 0:
                label = report_def.totals_label
                if "count" in req.totals:
                    label = f"{label} ({format_value(req.totals['count'], ColType.INTEGER, req.currency)} assets)"
                _add_run(p, label, size_pt=8, bold=True, colour=_C_WHITE)
            elif col_def.key in req.totals:
                font_name = "Montserrat" if col_def.col_type in (
                    ColType.CURRENCY, ColType.NUMBER, ColType.INTEGER
                ) else "Montserrat"
                _add_run(p, format_value(req.totals[col_def.key], col_def.col_type, req.currency),
                         font_name=font_name, size_pt=8, bold=True, colour=_C_WHITE)

        _set_row_height(totals_row, 7.0)

    # Style table borders
    _style_data_table_borders(table)

    # ── Authorisation block ───────────────────────────────────────────────
    doc.add_paragraph().space_after = Pt(12)

    p_auth_title = doc.add_paragraph()
    p_auth_title.space_before = Pt(8)
    p_auth_title.space_after = Pt(4)
    run_auth = _add_run(p_auth_title, "Authorisation", size_pt=12, bold=True, colour=_C_VIOLET)

    # Underline
    pPr = p_auth_title._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="615AA0"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    auth_table = doc.add_table(rows=3, cols=2)
    auth_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for col_idx, label in enumerate(["Prepared by", "Approved by"]):
        cell = auth_table.cell(0, col_idx)
        p = cell.paragraphs[0]
        _add_run(p, label, size_pt=8, colour=_C_GREY)

        # Signature line (empty row with bottom border)
        sig_cell = auth_table.cell(1, col_idx)
        sig_cell.text = ""
        _set_row_height(auth_table.rows[1], 14.0)

        # Name / Date label
        date_cell = auth_table.cell(2, col_idx)
        p_date = date_cell.paragraphs[0]
        _add_run(p_date, "Name / Date", size_pt=7, colour=_C_GREY)

    # ── Footer ────────────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_footer.space_before = Pt(4)

    footer_text = "Confidential"
    if req.company.legal_name:
        footer_text += f" · {req.company.legal_name}"
    footer_text += " · This report was generated by the Vairiot platform."

    _add_run(p_footer, footer_text, size_pt=6, colour=_C_GREY)

    # Page numbers (right-aligned — separate paragraph)
    p_page = footer.add_paragraph()
    p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p_page, "Page ", size_pt=6, colour=_C_GREY)
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_page = p_page.add_run()
    run_page.font.size = Pt(6)
    run_page.font.color.rgb = _C_GREY
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    _add_run(p_page, " of ", size_pt=6, colour=_C_GREY)
    # Total pages
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_total = p_page.add_run()
    run_total.font.size = Pt(6)
    run_total.font.color.rgb = _C_GREY
    run_total._r.append(fldChar3)
    run_total._r.append(instrText2)
    run_total._r.append(fldChar4)

    # Bottom gradient accent bar
    p_spacer = doc.add_paragraph()
    p_spacer.space_before = Pt(4)
    p_spacer.space_after = Pt(0)
    accent_bottom = doc.add_table(rows=1, cols=gradient_steps)
    accent_bottom.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, colour in enumerate(gradient_colours):
        cell = accent_bottom.cell(0, i)
        _set_cell_shading(cell, colour)
        cell.text = ""
        cell.paragraphs[0].space_before = Pt(0)
        cell.paragraphs[0].space_after = Pt(0)
    _set_row_height(accent_bottom.rows[0], 1.5)
    _remove_table_borders(accent_bottom)

    # ── Output ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _style_data_table_borders(table):
    """Style the data table with thin grey horizontal borders only."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D9D9D9"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="D9D9D9"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
